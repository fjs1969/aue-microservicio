from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from .pdf_processor import procesar_pdfs  # Asegúrate de que este archivo exista

app = FastAPI()

class InputData(BaseModel):
    municipio: str
    url_ficha: str
    url_informe: str

@app.post("/procesar")
def procesar(data: InputData):
    try:
        # Log de entrada para depuración
        print(f"📥 Recibido: municipio='{data.municipio}' url_ficha='{data.url_ficha}'; url_informe='{data.url_informe}'")

        # Procesar PDFs y generar contenido
        resultado = procesar_pdfs(data.municipio, data.url_ficha, data.url_informe)

        # Crear archivo Word con el resultado
        doc = Document()
        doc.add_heading(f"Diagnóstico AUE - {data.municipio}", 0)
        doc.add_paragraph(resultado)

        # Guardar con nombre normalizado
        nombre_docx = f"{data.municipio.lower()}_diagnostico.docx".replace(" ", "_")
        doc.save(nombre_docx)

        # Devolver el archivo como descarga
        return FileResponse(
            path=nombre_docx,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=nombre_docx
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
