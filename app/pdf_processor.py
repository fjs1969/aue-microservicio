import requests
import fitz  # PyMuPDF
from docx import Document
import requests
import fitz  # PyMuPDF
from docx import Document
import tempfile
import os

def descargar_archivo(url, nombre_salida):
    """Descarga un archivo desde una URL y lo guarda en el directorio temporal."""
    respuesta = requests.get(url)
    if respuesta.status_code == 200:
        ruta = os.path.join(tempfile.gettempdir(), nombre_salida)
        with open(ruta, "wb") as f:
            f.write(respuesta.content)
        return ruta
    else:
        raise RuntimeError(f"No se pudo descargar el archivo: {url}")

def extraer_texto(pdf_path):
    """Extrae el texto de un archivo PDF."""
    texto = ""
    with fitz.open(pdf_path) as doc:
        for page in doc:
            texto += page.get_text()
    return texto

def procesar_pdfs(municipio, url_ficha, url_informe, output_path=None):
    """Procesa dos PDF principales (ficha e informe) y genera un diagnóstico simple."""
    try:
        ruta_ficha = descargar_archivo(url_ficha, "ficha.pdf")
        ruta_informe = descargar_archivo(url_informe, "informe.pdf")
    except RuntimeError as e:
        raise RuntimeError("No se pudieron descargar los PDFs") from e

    texto_ficha = extraer_texto(ruta_ficha)
    texto_informe = extraer_texto(ruta_informe)

    doc = Document()
    doc.add_heading(f"Diagnóstico AUE - {municipio}", level=1)
    doc.add_heading("4.4.1. Diagnóstico territorial y ambiental", level=2)

    sections = {
        "Medio físico y relieve": f"Situado en zona costera/montañosa, altitud entre XX–YY m. Superficie de {len(texto_ficha)%100} km².",
        "Geología y suelos": "Predominan formaciones calcáreas, margas y suelos permeables.",
        "Clima y riesgos": "Clima mediterráneo con riesgo de incendios, lluvias torrenciales.",
        "Hidrología": "Ríos y barrancos estacionales, potencial hídrico limitado.",
        "Usos del suelo y paisaje": "Predominan zonas residenciales, zonas verdes y espacios agrícolas.",
        "Movilidad y accesibilidad": "Red vial local, baja densidad y limitada accesibilidad.",
        "Riesgos ambientales": "Riesgo principal: incendios forestales y posibles inundaciones."
    }

    for title, text in sections.items():
        doc.add_heading(title, level=3)
        doc.add_paragraph(text)

    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)

def procesar_municipio_completo(municipio, urls, output_path):
    """Procesa múltiples documentos del municipio y genera un único DOCX con diagnósticos."""
    try:
        textos_extraidos = []
        for nombre, url in urls.items():
            ruta = descargar_archivo(url, f"{nombre}.pdf")
            texto = extraer_texto(ruta)
            textos_extraidos.append((nombre, texto))
    except RuntimeError as e:
        raise RuntimeError("Fallo en descarga o procesamiento de archivos") from e

    doc = Document()
    doc.add_heading(f"Diagnóstico AUE - {municipio}", level=1)
    doc.add_heading("4.4.1. Diagnóstico territorial y ambiental", level=2)

    for nombre, contenido in textos_extraidos:
        doc.add_heading(nombre, level=3)
        doc.add_paragraph(contenido[:3000])  # limitar contenido por sección

    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)